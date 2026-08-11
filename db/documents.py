"""Document processing: parse, chunk, embed, and retrieve .doc/.docx/.pdf/.txt files.

Flow:
  1. Upload .doc/.docx/.pdf/.txt/.md file
  2. Extract plain text (parser per file type)
  3. Split into overlapping chunks
  4. Embed chunks with a local multilingual model (CPU)
  5. Store in ChromaDB for semantic search (RAG)

At query time, relevant chunks are injected into the LLM prompt so the model
can answer using the uploaded documents.
"""

import shutil
import subprocess
import uuid
from pathlib import Path
from typing import Optional

from loguru import logger

from config import settings


# ═══════════════════════════════════════════════════════════════════════════════
# Text extraction
# ═══════════════════════════════════════════════════════════════════════════════

def extract_text(filepath: Path) -> str:
    """Extract plain text from a document based on its extension.

    Supports: .docx, .doc, .pdf, .txt, .md

    Args:
        filepath: Path to the document.

    Returns:
        Extracted text.

    Raises:
        ValueError: If the file type is unsupported or extraction fails.
    """
    ext = filepath.suffix.lower()

    if ext in (".txt", ".md", ".markdown", ".rtf", ".log"):
        return _extract_plain(filepath)
    if ext == ".docx":
        return _extract_docx(filepath)
    if ext == ".doc":
        return _extract_doc(filepath)
    if ext == ".pdf":
        return _extract_pdf(filepath)
    raise ValueError(f"Unsupported file type: {ext} (supported: .doc, .docx, .pdf, .txt, .md)")


def _extract_plain(filepath: Path) -> str:
    """Plain text file — just read it."""
    for enc in ("utf-8", "cp1251", "latin-1"):
        try:
            return filepath.read_text(encoding=enc)
        except UnicodeDecodeError:
            continue
    # Last resort: read with replacement
    return filepath.read_text(encoding="utf-8", errors="replace")


def _extract_docx(filepath: Path) -> str:
    """Extract text from .docx using python-docx."""
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ValueError("python-docx not installed. Run: pip install python-docx")

    doc = DocxDocument(str(filepath))
    parts: list[str] = []

    # Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _extract_doc(filepath: Path) -> str:
    """Extract text from legacy .doc (binary Word format).

    Tries, in order:
      1. antiword   (best for .doc)
      2. catdoc
      3. LibreOffice headless conversion to txt
      4. python-docx fallback (works for some .doc that are actually .docx)
    """
    # 1. antiword
    if shutil.which("antiword"):
        try:
            result = subprocess.run(
                ["antiword", "-w", "0", str(filepath)],
                capture_output=True, text=True, timeout=120,
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout
        except (subprocess.TimeoutExpired, OSError):
            logger.warning("antiword failed, trying catdoc...")

    # 2. catdoc
    if shutil.which("catdoc"):
        try:
            result = subprocess.run(
                ["catdoc", str(filepath)],
                capture_output=True, text=True, timeout=120,
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout
        except (subprocess.TimeoutExpired, OSError):
            logger.warning("catdoc failed, trying LibreOffice...")

    # 3. LibreOffice headless
    if shutil.which("soffice") or shutil.which("libreoffice"):
        soffice = shutil.which("soffice") or shutil.which("libreoffice")
        import tempfile
        with tempfile.TemporaryDirectory() as tmp:
            result = subprocess.run(
                [soffice, "--headless", "--convert-to", "txt:Text", "--outdir", tmp, str(filepath)],
                capture_output=True, text=True, timeout=180,
            )
            if result.returncode == 0:
                out_name = filepath.stem + ".txt"
                out_file = Path(tmp) / out_name
                if out_file.exists():
                    return _extract_plain(out_file)

    # 4. Fallback: some ".doc" files are actually docx
    try:
        return _extract_docx(filepath)
    except Exception:
        pass

    raise ValueError(
        "Cannot extract text from .doc file. Install: "
        "sudo apt install antiword catdoc"
    )


def _extract_pdf(filepath: Path) -> str:
    """Extract text from .pdf using pypdf."""
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ValueError("pypdf not installed. Run: pip install pypdf")

    reader = PdfReader(str(filepath))
    parts = []
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
            if text.strip():
                parts.append(f"[Стр. {i+1}]\n{text}")
        except Exception:
            continue
    return "\n\n".join(parts)


# ═══════════════════════════════════════════════════════════════════════════════
# Chunking
# ═══════════════════════════════════════════════════════════════════════════════

def chunk_text(
    text: str,
    chunk_size: Optional[int] = None,
    overlap: Optional[int] = None,
) -> list[str]:
    """Split text into overlapping chunks.

    Splits on paragraph boundaries first, then sentence/word boundaries,
    so chunks are coherent (not cutting mid-word).

    Args:
        text: Full document text.
        chunk_size: Max characters per chunk.
        overlap: Characters of overlap between chunks.

    Returns:
        List of chunk strings.
    """
    chunk_size = chunk_size or settings.chunk_size
    overlap = overlap or settings.chunk_overlap

    if not text.strip():
        return []

    # Normalize whitespace
    text = "\n".join(line.rstrip() for line in text.splitlines())
    text = text.strip()

    # Split into paragraphs
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    if not paragraphs:
        paragraphs = [text]

    chunks: list[str] = []
    current = ""

    for para in paragraphs:
        # If a single paragraph is longer than chunk_size, hard-split it
        while len(para) > chunk_size:
            # Find a good split point (space) near chunk_size
            cut = chunk_size
            if current:
                cut = chunk_size - len(current) - 1
            split_at = para.rfind(" ", 0, cut)
            if split_at < chunk_size // 2:
                split_at = cut  # no good split point, hard cut

            piece = para[:split_at].strip()
            para = para[split_at:].strip()

            if current:
                current = f"{current} {piece}"
            else:
                current = piece

            if len(current) >= chunk_size:
                chunks.append(current)
                # overlap tail
                tail = current[-overlap:]
                current = tail

        # Append paragraph to current chunk
        if current and len(current) + len(para) + 1 > chunk_size:
            chunks.append(current)
            current = current[-overlap:]  # carry overlap
            current = f"{current} {para}".strip() if current else para
        else:
            current = f"{current} {para}".strip() if current else para

    if current.strip():
        chunks.append(current)

    # Filter empty chunks
    return [c for c in chunks if c.strip()]


# ═══════════════════════════════════════════════════════════════════════════════
# Document store: ChromaDB integration
# ═══════════════════════════════════════════════════════════════════════════════

DOCUMENTS_COLLECTION = "valera_documents"

# Lazy-loaded embedding function + collection (per-process singleton)
_embedding_fn = None
_collection = None


def _get_embedding_fn():
    """Lazily load the embedding function (local, CPU)."""
    global _embedding_fn
    if _embedding_fn is None:
        from chromadb.utils import embedding_functions
        _embedding_fn = embedding_functions.SentenceTransformerEmbeddingFunction(
            model_name=settings.embedding_model,
            device=settings.embedding_device,
        )
        logger.info(f"Embedding model loaded: {settings.embedding_model} (CPU)")
    return _embedding_fn


def _get_collection():
    """Lazily load/create the ChromaDB collection for documents."""
    global _collection
    if _collection is None:
        import chromadb
        from chromadb.config import Settings as ChromaSettings

        client = chromadb.PersistentClient(
            path=str(settings.chroma_path),
            settings=ChromaSettings(anonymized_telemetry=False),
        )
        try:
            _collection = client.get_collection(
                name=DOCUMENTS_COLLECTION,
                embedding_function=_get_embedding_fn(),
            )
        except Exception:
            _collection = client.create_collection(
                name=DOCUMENTS_COLLECTION,
                embedding_function=_get_embedding_fn(),
                metadata={"hnsw:space": "cosine"},
            )
            logger.info(f"Created document collection: {DOCUMENTS_COLLECTION}")
    return _collection


def index_document(doc_id: int, filename: str, text: str) -> int:
    """Chunk and embed a document's text into ChromaDB.

    Args:
        doc_id: Document ID in SQLite.
        filename: Original filename (for metadata).
        text: Extracted text.

    Returns:
        Number of chunks stored.
    """
    chunks = chunk_text(text)
    if not chunks:
        logger.warning(f"Document '{filename}' produced no chunks.")
        return 0

    collection = _get_collection()
    ids = [f"doc_{doc_id}_chunk_{i}" for i in range(len(chunks))]
    metadatas = [
        {
            "doc_id": str(doc_id),
            "filename": filename,
            "chunk_index": i,
        }
        for i in range(len(chunks))
    ]

    # Batch to avoid memory spikes
    batch_size = 100
    for start in range(0, len(chunks), batch_size):
        end = min(start + batch_size, len(chunks))
        collection.add(
            documents=chunks[start:end],
            metadatas=metadatas[start:end],
            ids=ids[start:end],
        )

    logger.info(f"Indexed {filename}: {len(chunks)} chunks stored.")
    return len(chunks)


def delete_document_chunks(doc_id: int) -> int:
    """Remove all chunks belonging to a document from ChromaDB.

    Returns:
        Number of chunks deleted (approx).
    """
    try:
        collection = _get_collection()
        # Get ids for this document
        result = collection.get(
            where={"doc_id": str(doc_id)},
            include=[],
        )
        ids = result.get("ids", [])
        if ids:
            collection.delete(ids=ids)
            logger.info(f"Deleted {len(ids)} chunks for document {doc_id}")
            return len(ids)
    except Exception as e:
        logger.warning(f"Could not delete chunks for doc {doc_id}: {e}")
    return 0


def search_documents(query: str, top_k: Optional[int] = None) -> list[dict]:
    """Semantic search over indexed documents.

    Args:
        query: User question / search query.
        top_k: Number of results.

    Returns:
        List of dicts: {text, filename, doc_id, chunk_index, distance}
    """
    top_k = top_k or settings.rag_top_k

    try:
        collection = _get_collection()
        results = collection.query(
            query_texts=[query],
            n_results=top_k,
        )
    except Exception as e:
        logger.error(f"Document search failed: {e}")
        return []

    formatted = []
    if results["documents"] and results["documents"][0]:
        for i in range(len(results["documents"][0])):
            meta = (
                results["metadatas"][0][i]
                if results["metadatas"] and results["metadatas"][0]
                else {}
            )
            formatted.append({
                "text": results["documents"][0][i],
                "filename": meta.get("filename", "unknown"),
                "doc_id": meta.get("doc_id", "0"),
                "chunk_index": meta.get("chunk_index", 0),
                "distance": (
                    results["distances"][0][i]
                    if results["distances"] and results["distances"][0]
                    else None
                ),
            })
    return formatted


def search_documents_formatted(query: str, top_k: Optional[int] = None) -> str:
    """Search documents and return a formatted block for the LLM prompt."""
    results = search_documents(query, top_k)
    if not results:
        return ""

    lines = ["Информация из загруженных документов:"]
    for i, r in enumerate(results, 1):
        text_short = r["text"][:600] + "..." if len(r["text"]) > 600 else r["text"]
        lines.append(f"[Документ {i}: {r['filename']}]\n{text_short}")

    return "\n\n".join(lines)


def get_document_count() -> int:
    """Number of chunks in the document collection."""
    try:
        return _get_collection().count()
    except Exception:
        return 0


def reset_embedding_cache():
    """Clear the lazy-loaded embedding singleton (e.g., after config change)."""
    global _embedding_fn, _collection
    _embedding_fn = None
    _collection = None

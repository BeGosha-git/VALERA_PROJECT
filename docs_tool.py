"""CLI tool for managing the document knowledge base.

Lets you index, search, list and delete documents WITHOUT starting the server.

Usage:
    python docs_tool.py index path/to/file.docx
    python docs_tool.py index path/to/folder/          # index all docs in folder
    python docs_tool.py search "мой вопрос"
    python docs_tool.py list
    python docs_tool.py stats
    python docs_tool.py delete <doc_id>
    python docs_tool.py test                            # create & index a sample .docx
"""

import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))

from loguru import logger
from config import settings
from db.database import db
from db.documents import (
    delete_document_chunks,
    extract_text,
    get_document_count,
    index_document,
    search_documents,
)

logger.remove()
logger.add(sys.stderr, level="INFO")

ALLOWED = {".doc", ".docx", ".pdf", ".txt", ".md", ".markdown", ".rtf", ".log"}


def cmd_index(paths: list[str]):
    """Index one or more files/folders."""
    settings.ensure_dirs()
    files: list[Path] = []
    for p in paths:
        path = Path(p)
        if path.is_dir():
            files.extend(
                f for f in path.rglob("*")
                if f.is_file() and f.suffix.lower() in ALLOWED
            )
        elif path.is_file() and path.suffix.lower() in ALLOWED:
            files.append(path)
        else:
            print(f"⚠️  Skipping (unsupported or not found): {p}")

    if not files:
        print("❌ No supported files found.")
        return

    for f in files:
        print(f"\n📄 Indexing: {f.name}")
        try:
            # Register in DB
            doc = db.add_document(
                filename=f.name,
                stored_path=str(f),
                file_type=f.suffix.lower(),
                file_size=f.stat().st_size,
            )
            # Extract text
            text = extract_text(f)
            if not text.strip():
                db.update_document(doc.id, status="error", error="Empty text")
                print("  ⚠️  No text extracted")
                continue
            # Chunk + embed
            n = index_document(doc.id, f.name, text)
            db.update_document(doc.id, num_chunks=n, status="ready")
            print(f"  ✅ {n} chunks indexed")
        except Exception as e:
            print(f"  ❌ Error: {e}")


def cmd_search(query: str, top_k: int = 5):
    """Search documents semantically."""
    results = search_documents(query, top_k)
    if not results:
        print("❌ Nothing found.")
        return
    print(f"\n🔍 Results for: '{query}' (top {top_k}):\n")
    for i, r in enumerate(results, 1):
        print(f"[{i}] 📄 {r['filename']} (chunk {r['chunk_index']})")
        print(f"    {r['text'][:300].replace(chr(10), ' ')}...")
        print()


def cmd_list():
    """List all documents."""
    docs = db.get_all_documents()
    if not docs:
        print("No documents yet.")
        return
    print(f"\n📚 Documents ({len(docs)}):\n")
    for d in docs:
        status_icon = {"ready": "✅", "processing": "⏳", "error": "❌"}.get(d.status, "?")
        print(f"  [{d.id}] {status_icon} {d.filename} — {d.num_chunks} chunks ({d.status})")


def cmd_stats():
    """Show stats."""
    docs = db.get_all_documents()
    print(f"\n📊 Document KB stats:")
    print(f"  Documents: {len(docs)}")
    print(f"  Total chunks: {get_document_count()}")
    print(f"  Ready: {sum(1 for d in docs if d.status == 'ready')}")
    print(f"  Processing: {sum(1 for d in docs if d.status == 'processing')}")
    print(f"  Errors: {sum(1 for d in docs if d.status == 'error')}")


def cmd_delete(doc_id: int):
    """Delete a document and its chunks."""
    doc = db.get_document(doc_id)
    if not doc:
        print(f"❌ Document {doc_id} not found.")
        return
    deleted = delete_document_chunks(doc_id)
    db.delete_document(doc_id)
    print(f"🗑️  Deleted '{doc.filename}' + {deleted} chunks.")


def cmd_test():
    """Create a sample .docx and index it (tests the whole pipeline)."""
    print("Creating sample .docx...")
    try:
        from docx import Document as DocxDocument
    except ImportError:
        print("❌ python-docx not installed. Run: pip install python-docx")
        return

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "sample.docx"
        doc = DocxDocument()
        doc.add_heading("Мой тестовый документ", level=1)
        doc.add_paragraph(
            "Qwen3-Omni — это end-to-end мультимодальная модель от Alibaba. "
            "Она понимает текст, аудио, изображения и видео, и генерирует "
            "речь. В этом проекте она используется как голосовой ассистент."
        )
        doc.add_paragraph(
            "База знаний хранит загруженные документы. Каждый документ "
            "разбивается на чанки, которые превращаются в векторные "
            "эмбеддинги для семантического поиска. При запросе нужные "
            "чанки автоматически подставляются в промпт модели."
        )
        doc.add_paragraph(
            "Интернет-поиск использует DuckDuckGo, а внутренняя база "
            "данных — SQLite и ChromaDB. Всё работает локально."
        )
        doc.save(str(path))

        print(f"  Sample saved: {path}")
        cmd_index([str(path)])

        print("\n🔎 Searching for 'что такое база знаний'...")
        cmd_search("что такое база знаний", 2)


def main():
    if len(sys.argv) < 2:
        print(__doc__)
        return

    cmd = sys.argv[1]

    if cmd == "index" and len(sys.argv) >= 3:
        cmd_index(sys.argv[2:])
    elif cmd == "search" and len(sys.argv) >= 3:
        cmd_search(sys.argv[2], int(sys.argv[3]) if len(sys.argv) > 3 else 5)
    elif cmd == "list":
        cmd_list()
    elif cmd == "stats":
        cmd_stats()
    elif cmd == "delete" and len(sys.argv) >= 3:
        cmd_delete(int(sys.argv[2]))
    elif cmd == "test":
        cmd_test()
    else:
        print(__doc__)


if __name__ == "__main__":
    main()
